import docx

try:
    # Open the Word document
    doc = docx.Document('Data Science Take-Home Assignment.docx')
    
    # Print the content paragraph by paragraph
    print("Document Content:\n")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():  # Check if paragraph has content
            print(f"Paragraph {i+1}: {para.text}")
            print("-" * 80)
    
    # Print information about tables if any
    if doc.tables:
        print(f"\nThe document contains {len(doc.tables)} tables.")
        
except Exception as e:
    print(f"Error reading Word document: {e}")